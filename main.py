import configparser
from os import PathLike
import sys

from docx import Document as DocxDocument
from datetime import datetime, timedelta
import os
import subprocess
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
import pickle
import json
import base64


def replace_text_in_docx(doc_path, replacements):
    doc = DocxDocument(doc_path)

    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                paragraph.text = paragraph.text.replace(old_text, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in paragraph.text:
                            paragraph.text = paragraph.text.replace(old_text, new_text)

    return doc


def convert_to_pdf(docx_path, pdf_path):
    libreoffice_path = "/usr/bin/libreoffice"  # This is the default path on Ubuntu
    if not os.path.exists(libreoffice_path):
        libreoffice_path = "/usr/bin/soffice"  # Try alternate path

    if not os.path.exists(libreoffice_path):
        raise FileNotFoundError("LibreOffice not found. Please install it.")

    subprocess.run([
        libreoffice_path,
        '--headless',
        '--convert-to', 'pdf',
        '--outdir', os.path.dirname(pdf_path),
        docx_path
    ], check=True)

    # Rename the output file to match the desired pdf_path
    os.rename(os.path.splitext(docx_path)[0] + '.pdf', pdf_path)


def make_quittance(property_dict: dict, months_offset: int) -> str:
    address = dict(property_dict)["address"]
    tenantname = dict(property_dict)["tenantname"]
    landlordname = dict(property_dict)["landlordname"]
    hors_charge = dict(property_dict)["hors_charge"]
    charge = dict(property_dict)["charge"]
    total_litteral = dict(property_dict)["total_litteral"]
    rent_amt = int(hors_charge) + int(charge)
    date_format = "%d-%m-%Y"

    # Calculate dates
    today = datetime.today().replace(month=datetime.today().month - months_offset)  # Replace with today's date
    first_day_this_month = today.replace(day=1).strftime(date_format)
    first_day_next_month = (today.replace(day=1) + timedelta(days=32)).replace(day=1).strftime(date_format)
    thirteenth_this_month = today.replace(day=8).strftime(date_format)
    quittance_date = today.strftime(date_format)

    # Define replacements
    replacements = {
        "{address}": address,
        "{tenantname}": tenantname,
        "{rent_letters}": total_litteral,
        "{rent_amt}": str(rent_amt) + "€",
        "{start}": first_day_this_month,
        "{end}": first_day_next_month,
        "{ex_charge}": str(hors_charge) + "€",
        "{charge}": str(charge) + "€",
        "{recu}": thirteenth_this_month,
        "{signed}": quittance_date
    }
    output_doc = f'{property_name}_quittance_{today.strftime(date_format)}.docx'
    output_pdf = f'{property_name}_quittance_{today.strftime(date_format)}.pdf'
    modified_doc = replace_text_in_docx(input_doc, replacements)
    modified_doc.save(output_doc)
    print(f"Word document has been updated and saved as {output_doc}")
    # Convert to PDF
    convert_to_pdf(output_doc, output_pdf)
    print(f"PDF has been created as {output_pdf}")
    os.remove(output_doc)

    msg_body = f"Cher {tenantname},\n\nVeuillez trouver ci-joint votre quittance de loyer du {first_day_this_month} au {first_day_next_month}.\n\nMerci,\n{landlordname}"

    return msg_body, output_pdf

def get_gmail_credentials():
    SCOPES = ['https://www.googleapis.com/auth/gmail.send']
    creds = None
    
    # Load credentials from token.pickle if it exists
    if os.path.exists('config/token.pickle'):
        with open('config/token.pickle', 'rb') as token:
            creds = pickle.load(token)
    
    # If credentials don't exist or are invalid, get new ones
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            # Load client secrets from the config directory
            client_secret_file = next(f for f in os.listdir('config') if f.startswith('client_secret_'))
            flow = InstalledAppFlow.from_client_secrets_file(
                f'config/{client_secret_file}', SCOPES)
            creds = flow.run_local_server(port=0)
        
        # Save the credentials for the next run
        with open('config/token.pickle', 'wb') as token:
            pickle.dump(creds, token)
    
    return creds

def send_email(sender_email, receiver_email, subject, body, attachment_path):
    creds = get_gmail_credentials()
    
    # Create message container
    msg = MIMEMultipart()
    msg['From'] = f'"{sender_email}" <{sender_email}>'
    msg['To'] = f'"{receiver_email}" <{receiver_email}>'
    msg['Subject'] = subject

    msg.attach(MIMEText(body, 'plain'))

    # Attach file
    with open(attachment_path, "rb") as attachment:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(attachment.read())

    encoders.encode_base64(part)
    part.add_header(
        "Content-Disposition",
        f"attachment; filename= {os.path.basename(attachment_path)}",
    )
    msg.attach(part)

    # Create raw email
    raw_message = base64.urlsafe_b64encode(msg.as_bytes()).decode('utf-8')
    
    # Create Gmail API service
    service = build('gmail', 'v1', credentials=creds)
    
    try:
        # Send message
        message = service.users().messages().send(
            userId='me',
            body={'raw': raw_message}
        ).execute()
        print(f"Message Id: {message['id']}")
        return message
    except Exception as e:
        print(f"Error details: {str(e)}")
        raise e


if __name__ == "__main__":
    # Define your variables
    input_doc = "quittance_template.docx"
    ini_file = configparser.ConfigParser(interpolation=None)
    ini_file.read(os.path.join(os.sep, os.getcwd(), 'config', 'config.ini'))
    properties_list = ini_file.sections()

    # auth
    sender_email = dict(ini_file.items(section="gmail"))["sender_email"]
    properties_list.pop(properties_list.index("gmail"))

    months_offset = int(sys.argv[1])

    for property_name in properties_list:
        property_dict = dict(ini_file.items(section=property_name))
        msg_body, output_pdf = make_quittance(property_dict, months_offset=months_offset)

        # Clean up the email address (remove any comments)
        receiver_email = property_dict['tenant_email'].split('#')[0].strip()
        subject = f"Quittance de loyer {property_name}"

        try:
            send_email(sender_email, receiver_email, subject, msg_body, output_pdf)
            print("Email sent successfully")
        except Exception as e:
            print(f"Error sending email: {e}")
            print(f"Email not sent for {property_name}")